from docx.shared import Inches
import os
from docx import Document
from tkinter import filedialog

doc = Document()
files = os.listdir()
for file in files:
    with open(file, "r") as file1:
        a = file1.read()
        print(file)
        question=input("Enter Question: ")
        doc.add_heading(question)
        doc.add_paragraph(a)
        screenshot=filedialog.askopenfilename()
        doc.add_picture(screenshot,width=Inches(5.23),height=Inches(2.3))
name=input("Enter File Name: ")
extn=".docx"
filename=f"{name}{extn}"
doc.save(filename)
