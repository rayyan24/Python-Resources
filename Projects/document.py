from docx import Document
from docx.shared import Inches
doc=Document()
with open("123.txt","r") as file:
    a=file.read()
    doc.add_paragraph(a)
    doc.add_picture("pic1.png",width=Inches(5.23),height=Inches(2.3))
doc.save("123.docx")